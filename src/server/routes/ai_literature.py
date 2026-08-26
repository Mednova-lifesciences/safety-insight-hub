"""
AI-assisted local literature screening.

NAFDAC GVP requires weekly screening of local medical literature that is
rarely indexed internationally. The deterministic keyword engine on the
frontend (src/services/literature/screener.ts) does the first-pass
flagging; this endpoint gives the reviewer a structured clinical reading
of a flagged article — products, reactions, seriousness criteria and a
risk rationale. It never creates or decides anything by itself: its
output is labelled AI assist and a human decides whether the article
becomes a signal or a case.
"""
from __future__ import annotations

import io
import json
import logging
from typing import Literal, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel

from ..ai.client import AiNotConfiguredError, AiRequestError, is_ai_configured, structured_completion
from ..ai.prompts import LITERATURE_ANALYSIS_PROMPT, PROMPT_VERSION
from ..ai.schemas import AiLiteratureAnalysis
from ..dependencies import AuthenticatedUser, get_current_user

logger = logging.getLogger(__name__)
router = APIRouter()

# Keeps a single analysis call's token footprint bounded regardless of
# document length, mirroring the PSUR PDF budget. Extraction always
# returns whatever was read (truncated at this bound) so the frontend's
# deterministic keyword engine can still screen the text even when the
# AI layer is unavailable.
MAX_DOC_CHARS = 60_000


def _extract_pdf_text(raw: bytes) -> tuple[str, Optional[int]]:
    import pdfplumber

    parts: list[str] = []
    char_budget = MAX_DOC_CHARS
    total_pages: Optional[int] = None
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        total_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            if char_budget <= 0:
                break
            page_text = (page.extract_text() or "").strip()
            if not page_text:
                continue
            chunk = f"\n\n--- page {i + 1} ---\n{page_text[:char_budget]}"
            parts.append(chunk)
            char_budget -= len(chunk)
    return "".join(parts), total_pages


def _extract_docx_text(raw: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:MAX_DOC_CHARS]


def _extract_plain_text(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")[:MAX_DOC_CHARS]
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")[:MAX_DOC_CHARS]


class LiteratureAnalyzeRequest(BaseModel):
    title: str
    text: str


class LiteratureAnalysisOut(BaseModel):
    is_safety_relevant: bool
    products: list[str]
    reaction_terms: list[str]
    seriousness_criteria: list[str]
    risk_level: Literal["HIGH", "MODERATE", "LOW"]
    summary: str
    rationale: str


class LiteratureAnalyzeResponse(BaseModel):
    analysis: Optional[LiteratureAnalysisOut]
    ai_used: bool
    prompt_version: str
    model: Optional[str] = None
    error: Optional[str] = None


class LiteratureDocumentResponse(BaseModel):
    extracted_text: str
    truncated: bool = False
    pages_extracted: Optional[int] = None
    analysis: Optional[LiteratureAnalysisOut]
    ai_used: bool
    prompt_version: str
    model: Optional[str] = None
    error: Optional[str] = None


async def _run_analysis(
    title: str, text: str
) -> tuple[Optional[LiteratureAnalysisOut], bool, Optional[str], Optional[str]]:
    """Shared AI layer for both analyze endpoints. Returns
    (analysis, ai_used, model, error) — never raises."""
    try:
        payload = {"title": title.strip(), "text": text}
        completion = await structured_completion(
            system_prompt=LITERATURE_ANALYSIS_PROMPT,
            user_content=json.dumps(payload),
            max_output_tokens=900,
        )
        parsed = AiLiteratureAnalysis.model_validate(completion.data)
        return (
            LiteratureAnalysisOut(
                is_safety_relevant=parsed.is_safety_relevant,
                products=parsed.products,
                reaction_terms=parsed.reaction_terms,
                seriousness_criteria=parsed.seriousness_criteria,
                risk_level=parsed.risk_level,
                summary=parsed.summary,
                rationale=parsed.rationale,
            ),
            True,
            completion.model,
            None,
        )
    except AiNotConfiguredError as exc:
        logger.info("Literature analysis skipped: %s", exc)
        return None, False, None, str(exc)
    except AiRequestError as exc:
        logger.error("Literature analysis failed: %s", exc)
        return None, False, None, "AI literature analysis unavailable."
    except Exception as exc:  # malformed/unvalidatable model output
        logger.error("Literature analysis returned unusable output: %s", exc)
        return None, False, None, "AI literature analysis returned an unusable response."


@router.get("/status")
async def ai_literature_status():
    return {"configured": is_ai_configured()}


@router.post("/analyze", response_model=LiteratureAnalyzeResponse)
async def analyze_literature(
    request: LiteratureAnalyzeRequest,
    user: AuthenticatedUser = Depends(get_current_user),
):
    text = request.text.strip()
    if not text:
        return LiteratureAnalyzeResponse(analysis=None, ai_used=False, prompt_version=PROMPT_VERSION)

    analysis, ai_used, model, error = await _run_analysis(request.title, text)
    return LiteratureAnalyzeResponse(
        analysis=analysis,
        ai_used=ai_used,
        prompt_version=PROMPT_VERSION,
        model=model,
        error=error,
    )


@router.post("/analyze-document", response_model=LiteratureDocumentResponse)
async def analyze_literature_document(
    file: UploadFile = File(...),
    user: AuthenticatedUser = Depends(get_current_user),
):
    """Extract text from an uploaded article (PDF, Word, or plain text)
    and, when the AI layer is configured, analyse it in the same request.
    The extracted text is always returned — the frontend's deterministic
    keyword engine screens it regardless of whether the AI ran, so an
    unconfigured OPENAI_API_KEY degrades the AI layer only, never the
    screening itself."""
    raw = await file.read()
    name = (file.filename or "").lower()

    pages_extracted: Optional[int] = None
    try:
        if name.endswith(".pdf"):
            text, pages_extracted = _extract_pdf_text(raw)
        elif name.endswith(".docx"):
            text = _extract_docx_text(raw)
        elif name.endswith((".txt", ".md", ".csv", ".rtf")):
            text = _extract_plain_text(raw)
        else:
            raise HTTPException(
                status_code=415,
                detail="Unsupported file type — upload a PDF, Word (.docx) or plain-text (.txt/.md) document.",
            )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Document text extraction failed for %s: %s", file.filename, exc)
        return LiteratureDocumentResponse(
            extracted_text="",
            ai_used=False,
            prompt_version=PROMPT_VERSION,
            error="Could not extract text from this document.",
        )

    if not text.strip():
        return LiteratureDocumentResponse(
            extracted_text="",
            pages_extracted=pages_extracted,
            ai_used=False,
            prompt_version=PROMPT_VERSION,
            error="No extractable text found in this document (it may be a scanned image without a text layer).",
        )

    truncated = len(text) >= MAX_DOC_CHARS
    title = (file.filename or "Uploaded document").rsplit(".", 1)[0]
    analysis, ai_used, model, error = await _run_analysis(title, text)

    return LiteratureDocumentResponse(
        extracted_text=text,
        truncated=truncated,
        pages_extracted=pages_extracted,
        analysis=analysis,
        ai_used=ai_used,
        prompt_version=PROMPT_VERSION,
        model=model,
        error=error,
    )
